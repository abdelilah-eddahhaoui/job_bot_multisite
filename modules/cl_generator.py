# ------------------ modules/cl_generator.py ------------------
"""
Handles cover letter generation by filling blank-based .docx templates and converting to PDF.
"""
import os
import subprocess
import re
from docx import Document

from modules.prompts import COVER_LETTER_PROMPT
from modules.utils import llm_chat

# Paths to blank-based templates
TEMPLATE_EN_DOCX = "assets/templates/template_motivation.docx"


def _cleanup_filled_text(text: str) -> str:
    """
    Keep only the actual letter returned by the LLM:
    • Throw away any echoed prompt headers, rules, markdown, etc.
    • Make sure no bullet symbols (•) survive if the prompt forbids them.
    """
    # Find *all* occurrences of the salutation.
    salutation_re = re.compile(r"Dear\s+Hiring\s+Team.*", re.I)
    matches = list(salutation_re.finditer(text))

    # If we found more than one, take the **last** one.
    if matches:
        text = text[matches[-1].start():]

    # Remove any line that still contains “→” (rules) or leading bullets.
    cleaned_lines = []
    for ln in text.splitlines():
        if ("→" in ln) or ln.strip().startswith("•"):
            continue
        cleaned_lines.append(ln.rstrip())

    return "\n".join(cleaned_lines).strip()


def generate_cover_letter(job_title, company, location):
    """
    Load the appropriate blank-template, instruct LLaMA 3 to fill in blanks, return a Document.
    """
    # Choose template
    template_path = TEMPLATE_EN_DOCX
    try:
        tpl_doc = Document(template_path)
    except Exception as e:
        print(f"⚠️ Failed to load template '{template_path}': {e}")
        return None

    # Get raw template text (preserve blanks)
    template_lines = [p.text for p in tpl_doc.paragraphs]
    template_text = "\n".join(template_lines)

    # Prepare prompt
    prompt = COVER_LETTER_PROMPT.format(
    job_title=job_title,
    company=company,
    location=location,
    template_text=template_text
    )
    
    try:
        model = "gpt-3.5-turbo" if os.getenv("LLM_BACKEND")=="openai" else "llama3"
        resp = llm_chat(messages=[{"role": "user", "content": prompt}], model=model)
        filled = resp['message']['content'].strip()
    except Exception as e:
        print(f"⚠️ Letter filling failed: {e}")
        return None    
        
    # Cleanup
    filled = _cleanup_filled_text(filled)

    # Build new Document
    new_doc = Document()
    for line in filled.split("\n"):
        new_doc.add_paragraph(line)
    return new_doc


def save_to_pdf(job_info, letter_doc, folder):
    """
    Save a Document object to PDF in the given folder using LibreOffice headless.
    """
    os.makedirs(folder, exist_ok=True)
    temp_docx = os.path.join(folder, "Cover_Letter.docx")
    letter_doc.save(temp_docx)
    try:
        subprocess.run([
            'libreoffice', '--headless', '--convert-to', 'pdf', temp_docx,
            '--outdir', folder
        ], check=True, stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL)
    except Exception as e:
        print(f"⚠️ PDF conversion failed: {e}")
