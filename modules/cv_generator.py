# ── modules/cv_generator.py ───
"""
CV customiser for any .docx template.

What it does
------------
1. Copies the user’s base CV (.docx) to an output path.
2. Injects up to 5 new keywords right after the first paragraph that
   contains the word 'skills' (case insensitive).
3. Normalises fonts to Times New Roman 11 pt where unspecified.
4. Checks that the resulting PDF is single-page (requires LibreOffice).
   If still >1 page, it retries with only the first 3 keywords.
5. Does NOT delete any paragraphs
   (Optional flag `remove_paragraph_startswith` exists below.)

If LibreOffice, python-docx, or PyMuPDF aren’t installed,
it copies the doc unmodified.
"""
from __future__ import annotations

import os, shutil, subprocess, logging
from typing import List, Optional

try:
    from docx import Document
    from docx.shared import Pt
except ImportError:       # python-docx not installed
    Document = None
try:
    import fitz  # PyMuPDF
except ImportError:
    fitz = None

log = logging.getLogger(__name__)



#### MODIFY FONTS AND LABELS INPUTS IF NEEDED ####
FONT_NAME  = "Times New Roman"
FONT_SIZE  = Pt(11) if Document else None
LABEL_SKILLS  = "skills"
#### MODIFY FONTS AND LABELS INPUTS IF NEEDED ####



# ────────────────────────────────────────────────────────────────────
def is_single_page(docx_path: str) -> bool:
    """Convert DOCX → PDF and return True if the PDF has exactly one page."""
    if fitz is None:
        # Can’t check page count; assume OK
        return True

    pdf_path = docx_path.replace(".docx", ".pdf")
    try:
        subprocess.run(
            ["libreoffice", "--headless", "--convert-to", "pdf", docx_path, "--outdir", os.path.dirname(docx_path)],
            check=True, stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL,
        )
        return fitz.open(pdf_path).page_count == 1
    except Exception as e:
        log.warning("Single-page check failed: %s", e)
        return True


# ────────────────────────────────────────────────────────────────────
def insert_keywords_into_doc(
    template_cv: str,
    keywords: List[str],
    output_path: str,
    *,
    remove_paragraph_startswith: Optional[str] = None,
) -> None:
    """
    Parameters
    ----------
    template_cv : str
        Path to the user’s base CV (.docx).
    keywords : List[str]
        Ranked list of keywords extracted from the job description.
    output_path : str
        Target .docx path (will be overwritten).
    remove_paragraph_startswith : Optional[str], default None
        If provided (e.g. "hobbies"), *and* a paragraph starting with that
        lowercase text exists, it will be removed. Defaults to None (no deletion).
    """
    if Document is None:
        log.warning("python-docx not installed; copying CV without changes.")
        shutil.copyfile(template_cv, output_path)
        return

    if not os.path.exists(template_cv):
        raise FileNotFoundError(f"Base CV template not found: {template_cv}")

    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    shutil.copyfile(template_cv, output_path)
    _inject_keywords(output_path, keywords, remove_paragraph_startswith)

    # Retry with fewer keywords if CV still >1 page
    if keywords and not is_single_page(output_path):
        log.info("CV exceeds one page; retrying with 3 keywords.")
        _inject_keywords(output_path, keywords[:3], remove_paragraph_startswith, overwrite=True)
        if not is_single_page(output_path):
            log.warning("CV still exceeds 1 page after retry.")


# ────────────────────────────────────────────────────────────────────
def _inject_keywords(
    docx_path: str,
    keywords: List[str],
    remove_paragraph_startswith: Optional[str],
    *,
    overwrite: bool = False,
) -> None:
    """Internal helper that mutates the .docx in place."""
    doc = Document(docx_path)

    # Determine which keywords are new
    existing = {w.lower() for p in doc.paragraphs for w in p.text.split()}
    to_add   = [k for k in keywords if k.lower() not in existing][:5]

    if to_add:
        for p in doc.paragraphs:
            if LABEL_SKILLS in p.text.lower():
                if "Additional Skills:" in p.text:
                    break 
                run = p.add_run("\\nAdditional Skills: " + ", ".join(to_add))
                run = p.add_run("\nAdditional Skills: " + ", ".join(to_add))
                run.font.name = FONT_NAME
                run.font.size = FONT_SIZE
                break

    # Optionally remove a section (disabled by default)
    if remove_paragraph_startswith:
        target = remove_paragraph_startswith.lower()
        for p in reversed(doc.paragraphs):
            if p.text.strip().lower().startswith(target):
                p._element.getparent().remove(p._element)
                break

    # Normalise fonts
    for p in doc.paragraphs:
        for r in p.runs:
            if not r.font.name:
                r.font.name = FONT_NAME

    doc.save(docx_path if overwrite else docx_path)


# ────────────────────────────────────────────────────────────────────
def convert_to_pdf_libreoffice(input_path: str, output_dir: str) -> None:
    try:
        subprocess.run(
            ["libreoffice", "--headless", "--convert-to", "pdf", input_path, "--outdir", output_dir],
            check=True, stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL,
        )
    except Exception as e:
        log.warning("PDF conversion failed: %s", e)
# ────────────────────────────────────────────────────────────────────
